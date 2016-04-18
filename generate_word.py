# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import RGBColor
import subprocess
import time
import os

qq={"GRE":u"GRE备考群: 34304320","TOEFL":u"托福备考群： 460749662","SAT":u"SAT备考群： 477484091","IELTS":u"雅思备考群: 333975398","GMAT":u"GMT备考群： 467092077","kaoyan": u"考研备考群 : 206117896"}
def convert_pdf1(txtpath,type):
  txt_file=open(txtpath,"r")
  title=txt_file.readline()
  content=txt_file.read()
  document = Document("template/"+type+".docx")
  cursor_paragraph=document.paragraphs[-1]

  styles=document.styles
  head_style=styles.add_style('Heading_title', WD_STYLE_TYPE.PARAGRAPH)
  #head_style.alignment=1
  head_style.font.size=Pt(20)
  head=cursor_paragraph.insert_paragraph_before(title.decode("utf-8"),style="Heading_title")
  head.alignment=1
  cursor_paragraph.insert_paragraph_before(content.decode("utf-8"))
  qq_style=styles.add_style('qq', WD_STYLE_TYPE.PARAGRAPH)
  qq_style.font.color.rgb=RGBColor(200,10,10)
  qq_style.font.size=Pt(15)
  qq_txt=cursor_paragraph.insert_paragraph_before(u"\n\n"+qq[type],style="qq")
  qq_txt.alignment=2
  #qq_txt.style.font.color=
  document.save('tem.docx')

  #subprocess.popen("unoconv -f  pdf tem.docx")
  os.popen("rm tem.docx")
def convert_pdf2(title,content,path,filetype):
  print title
  document = Document("template/"+filetype+".docx")
  cursor_paragraph=document.paragraphs[-1]
  styles=document.styles
  head_style=styles.add_style('Heading_title', WD_STYLE_TYPE.PARAGRAPH)
  #head_style.alignment=1
  head_style.font.size=Pt(20)
  head=cursor_paragraph.insert_paragraph_before(title.decode("utf-8"),style="Heading_title")
  head.alignment=1
  content_text=cursor_paragraph.insert_paragraph_before(content.decode("utf-8"))
  content_text.alignment=0
  qq_style=styles.add_style('qq', WD_STYLE_TYPE.PARAGRAPH)
  qq_style.font.color.rgb=RGBColor(200,10,10)
  qq_style.font.size=Pt(15)
  qq_txt=cursor_paragraph.insert_paragraph_before(u"\n\n"+qq[filetype],style="qq")
  qq_txt.alignment=2
  #qq_txt.style.font.color=
  tem_docname=path+'/'+filetype+"/"+title+".docx"
  print "呵呵"
  print tem_docname
  document.save(tem_docname)
  print "lowriter --convert-to pdf "+title+".docx"
  #subprocess.call(["lowriter","--convert-to","pdf",tem_docname])
  #os.popen("lowriter --convert-to pdf "+title+".docx")
  #os.popen("rm tem.docx")
if __name__=="__main__":
  convert_pdf2(u"托福口语&写作备考：模版是把双刃剑".encode("utf-8"),"sdfsdf","pdf/2016-03-20/","GRE")
